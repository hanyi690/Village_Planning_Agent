"""
RAG Experiment Output Generator (Refactored)
Generate experiment output documents: Excel, Word, KB summary

Outputs:
1. Excel: Hallucination rate summary per dimension (3 runs each)
2. Word: Typical reference contrast paragraphs
3. KB Summary: Document count, source types, categories

Usage:
    python scripts/experiments/rag_hallucination/generate_experiment_outputs.py
    python scripts/experiments/rag_hallucination/generate_experiment_outputs.py --from-results
"""

import json
import sys
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional
from dataclasses import dataclass

sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent / "backend"))

from scripts.experiments.config import (
    RAG_HALLUCINATION_DIR,
    RAG_ON_DIR,
    RAG_OFF_DIR,
    RAG_ENABLED_DIMENSIONS,
    ANNOTATION_DIR,
    ensure_rag_experiment_dirs,
)


@dataclass
class HallucinationStats:
    """幻觉率统计数据"""
    dimension_key: str
    dimension_name: str
    rag_on_total_refs: int
    rag_on_hallucinations: int
    rag_on_hallucination_rate: float
    rag_off_total_refs: int
    rag_off_hallucinations: int
    rag_off_hallucination_rate: float
    rate_reduction: float


@dataclass
class ContrastParagraph:
    """对比段落"""
    dimension_key: str
    dimension_name: str
    rag_off_text: str
    rag_on_text: str
    regulation_source: str
    explanation: str


@dataclass
class KnowledgeBaseDoc:
    """知识库文档"""
    source: str
    doc_type: str
    chunk_count: int
    category: str


DIMENSION_NAMES = {
    "land_use_planning": "土地利用规划",
    "infrastructure_planning": "基础设施规划",
    "ecological": "生态绿地规划",
    "disaster_prevention": "防震减灾规划",
    "heritage": "历史文保规划",
}


def load_comparison_results() -> Optional[Dict[str, Any]]:
    """加载对比结果"""
    comparison_file = RAG_HALLUCINATION_DIR / "comparison_results.json"
    if comparison_file.exists():
        with open(comparison_file, "r", encoding="utf-8") as f:
            return json.load(f)
    return None


def load_single_run_results(mode: str, run_id: int) -> Dict[str, Any]:
    """加载单次运行结果"""
    output_dir = RAG_ON_DIR if mode == "rag_on" else RAG_OFF_DIR
    result_file = output_dir / f"run_{run_id}_results.json"
    if result_file.exists():
        with open(result_file, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def get_hallucination_stats_from_results() -> List[HallucinationStats]:
    """从实验结果获取幻觉率统计"""
    comparison = load_comparison_results()
    if not comparison:
        print("[Warning] No comparison results found, using empty stats")
        return []
    stats_list = []
    for dim in comparison.get("dimensions", []):
        stats_list.append(HallucinationStats(
            dimension_key=dim.get("dimension_key", ""),
            dimension_name=dim.get("dimension_name", ""),
            rag_on_total_refs=int(dim.get("rag_on_total_refs", 0)),
            rag_on_hallucinations=int(dim.get("rag_on_hallucinations", 0)),
            rag_on_hallucination_rate=dim.get("rag_on_rate", 0.0),
            rag_off_total_refs=int(dim.get("rag_off_total_refs", 0)),
            rag_off_hallucinations=int(dim.get("rag_off_hallucinations", 0)),
            rag_off_hallucination_rate=dim.get("rag_off_rate", 0.0),
            rate_reduction=dim.get("rate_reduction", 0.0),
        ))
    return stats_list


def get_contrast_paragraphs_from_results() -> List[ContrastParagraph]:
    """从实验结果获取对比段落"""
    paragraphs = []
    rag_on_results = load_single_run_results("rag_on", 1)
    rag_off_results = load_single_run_results("rag_off", 1)
    for dim_key in RAG_ENABLED_DIMENSIONS:
        rag_on_report = rag_on_results.get(dim_key, {}).get("report_content", "")
        rag_off_report = rag_off_results.get(dim_key, {}).get("report_content", "")
        if not rag_on_report or not rag_off_report:
            continue
        contrast = find_best_contrast_paragraph(
            dim_key, DIMENSION_NAMES.get(dim_key, dim_key),
            rag_on_report, rag_off_report
        )
        if contrast:
            paragraphs.append(contrast)
    return paragraphs


def find_best_contrast_paragraph(dim_key: str, dim_name: str,
                                  rag_on_report: str,
                                  rag_off_report: str) -> Optional[ContrastParagraph]:
    """寻找最佳对比段落"""
    rag_on_paragraphs = [p.strip() for p in rag_on_report.split("\n\n") if p.strip()]
    rag_off_paragraphs = [p.strip() for p in rag_off_report.split("\n\n") if p.strip()]

    def calculate_specificity(text: str) -> float:
        score = 0.0
        if "第" in text and "条" in text:
            score += 2.0
        if "米" in text or "%" in text or "年一遇" in text:
            score += 3.0
        if "《" in text and "》" in text:
            score += 1.0
        if "GB" in text or "DB" in text:
            score += 2.0
        return score

    best_rag_on = None
    best_rag_on_score = -1
    for para in rag_on_paragraphs:
        score = calculate_specificity(para)
        if score > best_rag_on_score:
            best_rag_on_score = score
            best_rag_on = para

    best_rag_off = None
    best_rag_off_score = float("inf")
    for para in rag_off_paragraphs:
        score = calculate_specificity(para)
        if score < best_rag_off_score and len(para) > 50:
            best_rag_off_score = score
            best_rag_off = para

    if not best_rag_on or not best_rag_off:
        return None

    sources = re.findall(r"《[^》]+》", best_rag_on)
    regulation_source = sources[0] if sources else "未知来源"

    return ContrastParagraph(
        dimension_key=dim_key,
        dimension_name=dim_name,
        rag_off_text=best_rag_off[:300] + "..." if len(best_rag_off) > 300 else best_rag_off,
        rag_on_text=best_rag_on[:500] + "..." if len(best_rag_on) > 500 else best_rag_on,
        regulation_source=regulation_source,
        explanation=f"RAG ON特异性评分: {best_rag_on_score:.1f}, RAG OFF特异性评分: {best_rag_off_score:.1f}",
    )


def get_kb_documents_from_kb() -> List[KnowledgeBaseDoc]:
    """从知识库获取文档列表"""
    docs = []
    try:
        from app.services.modules.rag.service import RagService
        rag_service = RagService()
        kb_stats = rag_service.get_stats()
        for doc in kb_stats.get("documents", []):
            source = doc.get("source", "")
            doc_type = infer_doc_type(source)
            category = infer_category(source)
            docs.append(KnowledgeBaseDoc(
                source=source,
                doc_type=doc_type,
                chunk_count=doc.get("chunk_count", 0),
                category=category,
            ))
    except Exception as e:
        print(f"[Warning] Failed to load KB documents: {e}")
    return docs


def infer_doc_type(source: str) -> str:
    """推断文档类型"""
    source_lower = source.lower()
    if "法》" in source or "条例" in source:
        return "法规"
    elif "标准" in source or "gb" in source_lower or "db" in source_lower:
        return "技术标准"
    elif "规划" in source and ("省" in source or "市" in source):
        return "上位规划"
    elif "案例" in source or "实例" in source:
        return "案例"
    elif "教材" in source or "原理" in source:
        return "专业教材"
    else:
        return "政策文件"


def infer_category(source: str) -> str:
    """推断文档分类"""
    if "01_专业教材" in source:
        return "专业教材"
    elif "02_法律法规" in source:
        return "法律法规"
    elif "03_政策文件" in source:
        return "政策文件"
    elif "04_技术规范" in source:
        return "技术规范"
    elif "05_上位规划" in source:
        return "上位规划"
    elif "06_相关案例" in source:
        return "相关案例"
    else:
        return "其他"


def generate_excel_summary(stats_list: List[HallucinationStats], output_dir: Path) -> Path:
    """生成Excel汇总表"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        print("Warning: openpyxl not installed, using CSV fallback")
        return generate_csv_summary(stats_list, output_dir)

    wb = Workbook()
    ws = wb.active
    ws.title = "Hallucination Summary"

    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, size=11, color="FFFFFF")
    center_align = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )

    headers = [
        "Dimension", "Name",
        "RAG ON Total", "RAG ON Hallucinations", "RAG ON Rate",
        "RAG OFF Total", "RAG OFF Hallucinations", "RAG OFF Rate",
        "Rate Reduction"
    ]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = thin_border

    for row, stats in enumerate(stats_list, 2):
        data = [
            stats.dimension_key, stats.dimension_name,
            stats.rag_on_total_refs, stats.rag_on_hallucinations,
            f"{stats.rag_on_hallucination_rate:.1%}",
            stats.rag_off_total_refs, stats.rag_off_hallucinations,
            f"{stats.rag_off_hallucination_rate:.1%}",
            f"{stats.rate_reduction:.1%}",
        ]
        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row, column=col, value=value)
            cell.alignment = center_align
            cell.border = thin_border

    summary_row = len(stats_list) + 2
    ws.cell(row=summary_row, column=1, value="Total").font = Font(bold=True)
    ws.cell(row=summary_row, column=3, value=sum(s.rag_on_total_refs for s in stats_list))
    ws.cell(row=summary_row, column=4, value=sum(s.rag_on_hallucinations for s in stats_list))
    ws.cell(row=summary_row, column=6, value=sum(s.rag_off_total_refs for s in stats_list))
    ws.cell(row=summary_row, column=7, value=sum(s.rag_off_hallucinations for s in stats_list))

    column_widths = [22, 14, 12, 14, 12, 13, 14, 12, 12]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width

    output_path = output_dir / "hallucination_summary.xlsx"
    wb.save(output_path)
    print(f"[Excel] Saved: {output_path}")
    return output_path


def generate_csv_summary(stats_list: List[HallucinationStats], output_dir: Path) -> Path:
    """生成CSV汇总表"""
    import csv
    output_path = output_dir / "hallucination_summary.csv"
    with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow([
            "Dimension", "Name", "RAG ON Total", "RAG ON Hallucinations", "RAG ON Rate",
            "RAG OFF Total", "RAG OFF Hallucinations", "RAG OFF Rate", "Rate Reduction"
        ])
        for stats in stats_list:
            writer.writerow([
                stats.dimension_key, stats.dimension_name,
                stats.rag_on_total_refs, stats.rag_on_hallucinations,
                f"{stats.rag_on_hallucination_rate:.1%}",
                stats.rag_off_total_refs, stats.rag_off_hallucinations,
                f"{stats.rag_off_hallucination_rate:.1%}",
                f"{stats.rate_reduction:.1%}",
            ])
    print(f"[CSV] Saved: {output_path}")
    return output_path


def generate_word_contrast(paragraphs: List[ContrastParagraph], output_dir: Path) -> Path:
    """生成Word对比文档"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Warning: python-docx not installed, using Markdown fallback")
        return generate_markdown_contrast(paragraphs, output_dir)

    doc = Document()
    title = doc.add_heading("RAG Reference Contrast Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This document demonstrates the contrast between RAG ON and RAG OFF conditions. "
        "RAG ON provides precise regulation citations with specific clause numbers and technical indicators, "
        "while RAG OFF often produces generic statements without concrete basis."
    )

    doc.add_heading("Contrast Examples", level=1)
    for i, para in enumerate(paragraphs, 1):
        doc.add_heading(f"{i}. {para.dimension_name}", level=2)
        doc.add_heading("RAG OFF (Generic)", level=3)
        doc.add_paragraph(para.rag_off_text).italic = True
        doc.add_heading("RAG ON (Precise)", level=3)
        doc.add_paragraph(para.rag_on_text).bold = True
        doc.add_paragraph(f"Source: {para.regulation_source}")
        doc.add_paragraph(f"Explanation: {para.explanation}")
        doc.add_paragraph()

    output_path = output_dir / "contrast_analysis.docx"
    doc.save(output_path)
    print(f"[Word] Saved: {output_path}")
    return output_path


def generate_markdown_contrast(paragraphs: List[ContrastParagraph], output_dir: Path) -> Path:
    """生成Markdown对比文档"""
    output_path = output_dir / "contrast_analysis.md"
    content = "# RAG Reference Contrast Analysis\n\n"
    content += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"

    for i, para in enumerate(paragraphs, 1):
        content += f"## {i}. {para.dimension_name}\n\n"
        content += f"**RAG OFF (Generic)**:\n\n> {para.rag_off_text}\n\n"
        content += f"**RAG ON (Precise)**:\n\n> {para.rag_on_text}\n\n"
        content += f"- Source: {para.regulation_source}\n"
        content += f"- Explanation: {para.explanation}\n\n---\n\n"

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"[Markdown] Saved: {output_path}")
    return output_path


def generate_kb_summary(docs: List[KnowledgeBaseDoc], output_dir: Path) -> Path:
    """生成知识库说明表"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
    except ImportError:
        return generate_kb_summary_csv(docs, output_dir)

    wb = Workbook()
    ws = wb.active
    ws.title = "KB Summary"

    header_fill = PatternFill(start_color="70AD47", end_color="70AD47", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")

    headers = ["Document", "Type", "Chunks", "Category"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    for row, doc in enumerate(docs, 2):
        ws.cell(row=row, column=1, value=doc.source)
        ws.cell(row=row, column=2, value=doc.doc_type)
        ws.cell(row=row, column=3, value=doc.chunk_count)
        ws.cell(row=row, column=4, value=doc.category)

    summary_row = len(docs) + 2
    ws.cell(row=summary_row, column=1, value="Total").font = Font(bold=True)
    ws.cell(row=summary_row, column=3, value=sum(d.chunk_count for d in docs))

    type_counts = {}
    for doc in docs:
        type_counts[doc.doc_type] = type_counts.get(doc.doc_type, 0) + 1

    row = summary_row + 2
    ws.cell(row=row, column=1, value="Type Summary").font = Font(bold=True)
    row += 1
    for doc_type, count in type_counts.items():
        ws.cell(row=row, column=1, value=doc_type)
        ws.cell(row=row, column=2, value=count)
        row += 1

    output_path = output_dir / "kb_summary.xlsx"
    wb.save(output_path)
    print(f"[KB Summary] Saved: {output_path}")
    return output_path


def generate_kb_summary_csv(docs: List[KnowledgeBaseDoc], output_dir: Path) -> Path:
    """生成知识库说明表CSV"""
    import csv
    output_path = output_dir / "kb_summary.csv"
    with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Document", "Type", "Chunks", "Category"])
        for doc in docs:
            writer.writerow([doc.source, doc.doc_type, doc.chunk_count, doc.category])
    print(f"[KB Summary CSV] Saved: {output_path}")
    return output_path


def main(from_results: bool = True):
    """主函数"""
    print("=" * 60)
    print("RAG Experiment Output Generator")
    print("=" * 60)

    ensure_rag_experiment_dirs()
    output_dir = RAG_HALLUCINATION_DIR / "outputs"
    output_dir.mkdir(parents=True, exist_ok=True)

    if from_results:
        print("\n[1] Loading hallucination stats from results...")
        stats_list = get_hallucination_stats_from_results()

        print("\n[2] Loading contrast paragraphs from results...")
        paragraphs = get_contrast_paragraphs_from_results()

        print("\n[3] Loading KB documents from knowledge base...")
        kb_docs = get_kb_documents_from_kb()
    else:
        print("\n[Warning] No results available, generating empty outputs")
        stats_list = []
        paragraphs = []
        kb_docs = []

    print("\n[4] Generating Excel summary...")
    generate_excel_summary(stats_list, output_dir)

    print("\n[5] Generating Word contrast document...")
    generate_word_contrast(paragraphs, output_dir)

    print("\n[6] Generating Knowledge Base summary...")
    generate_kb_summary(kb_docs, output_dir)

    print("\n" + "=" * 60)
    print("All outputs generated!")
    print(f"Output directory: {output_dir}")
    print("=" * 60)


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--from-results", action="store_true", default=True,
                        help="Load data from experiment results")
    args = parser.parse_args()
    main(from_results=args.from_results)
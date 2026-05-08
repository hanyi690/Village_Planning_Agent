"""
分析参考 Word 文档结构

提取章节结构、段落模板、表格格式
"""

from docx import Document
from pathlib import Path
import json

# 参考文档路径
REF_DOC = Path(__file__).parent.parent.parent / "docs" / "平远县泗水镇金田村村庄规划（2022-2035年）文本.docx"


def analyze_document(doc_path: Path) -> dict:
    """分析 Word 文档结构"""
    doc = Document(doc_path)

    result = {
        "filename": doc_path.name,
        "paragraphs_count": len(doc.paragraphs),
        "tables_count": len(doc.tables),
        "sections": [],
        "paragraphs": [],
        "tables": [],
    }

    # 分析段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "None"
            is_heading = style and (style.startswith("Heading") or "标题" in style)
            para_info = {
                "index": i,
                "text_preview": text[:100] if len(text) > 100 else text,
                "full_text": text,
                "style": style or "Normal",
                "is_heading": is_heading,
            }
            result["paragraphs"].append(para_info)

            # 记录标题
            if is_heading:
                result["sections"].append({
                    "level": style,
                    "text": text,
                    "index": i,
                })

    # 分析表格
    for i, table in enumerate(doc.tables):
        table_info = {
            "index": i,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "headers": [],
            "sample_data": [],
        }

        # 提取表头（第一行）
        if table.rows:
            first_row = [cell.text.strip() for cell in table.rows[0].cells]
            table_info["headers"] = first_row

            # 提取前 3 行数据样本
            for row_idx in range(1, min(4, len(table.rows))):
                row_data = [cell.text.strip()[:50] for cell in table.rows[row_idx].cells]
                table_info["sample_data"].append(row_data)

        result["tables"].append(table_info)

    return result


def main():
    """主函数"""
    print(f"分析文档: {REF_DOC}")

    if not REF_DOC.exists():
        print(f"错误: 文档不存在 {REF_DOC}")
        return

    result = analyze_document(REF_DOC)

    # 输出摘要
    print(f"\n=== 文档摘要 ===")
    print(f"段落总数: {result['paragraphs_count']}")
    print(f"表格总数: {result['tables_count']}")
    print(f"标题数: {len(result['sections'])}")

    print(f"\n=== 章节结构 ===")
    for sec in result["sections"]:
        print(f"  [{sec['level']}] {sec['text']}")

    print(f"\n=== 表格结构 ===")
    for tbl in result["tables"]:
        print(f"  表格 {tbl['index']}: {tbl['rows']}行 x {tbl['cols']}列")
        print(f"    表头: {tbl['headers']}")
        if tbl['sample_data']:
            print(f"    数据样本: {tbl['sample_data'][0]}")

    # 保存完整分析结果
    output_file = Path(__file__).parent.parent.parent / "docs" / "reference_doc_structure.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    print(f"\n完整结构已保存到: {output_file}")


if __name__ == "__main__":
    main()
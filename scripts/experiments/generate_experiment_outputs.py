"""
RAG Experiment Output Generator
Generate experiment output documents: Excel, Word, KB summary

Outputs:
1. Excel: Hallucination rate summary per dimension
2. Word: Typical reference contrast paragraphs
3. KB Summary: Document count, source types, categories

Usage:
    python scripts/experiments/generate_experiment_outputs.py
    python scripts/experiments/generate_experiment_outputs.py --mock
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any
from dataclasses import dataclass

sys.path.insert(0, str(Path(__file__).parent.parent.parent / "backend"))

from scripts.experiments.config import (
    RAG_HALLUCINATION_DIR,
    RAG_ON_DIR,
    RAG_OFF_DIR,
    RAG_ENABLED_DIMENSIONS,
    ANNOTATION_DIR,
    ensure_rag_experiment_dirs,
)


@dataclass
class HallucinationStats:
    dimension_key: str
    dimension_name: str
    rag_on_total_refs: int
    rag_on_hallucinations: int
    rag_on_hallucination_rate: float
    rag_off_total_refs: int
    rag_off_hallucinations: int
    rag_off_hallucination_rate: float
    rate_reduction: float


@dataclass
class ContrastParagraph:
    dimension_key: str
    dimension_name: str
    rag_off_text: str
    rag_on_text: str
    regulation_source: str
    explanation: str


@dataclass
class KnowledgeBaseDoc:
    source: str
    doc_type: str
    chunk_count: int
    category: str


DIMENSION_NAMES = {
    "land_use_planning": "土地利用规划",
    "infrastructure_planning": "基础设施规划",
    "ecological": "生态绿地规划",
    "disaster_prevention": "防震减灾规划",
    "heritage": "历史文保规划",
}


def generate_excel_summary(stats_list: List[HallucinationStats], output_dir: Path) -> Path:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        print("Warning: openpyxl not installed, using CSV fallback")
        return generate_csv_summary(stats_list, output_dir)

    wb = Workbook()
    ws = wb.active
    ws.title = "Hallucination Summary"

    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, size=11, color="FFFFFF")
    center_align = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )

    headers = [
        "Dimension", "Name",
        "RAG ON Total", "RAG ON Hallucinations", "RAG ON Rate",
        "RAG OFF Total", "RAG OFF Hallucinations", "RAG OFF Rate",
        "Rate Reduction"
    ]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = thin_border

    for row, stats in enumerate(stats_list, 2):
        data = [
            stats.dimension_key, stats.dimension_name,
            stats.rag_on_total_refs, stats.rag_on_hallucinations,
            f"{stats.rag_on_hallucination_rate:.1%}",
            stats.rag_off_total_refs, stats.rag_off_hallucinations,
            f"{stats.rag_off_hallucination_rate:.1%}",
            f"{stats.rate_reduction:.1%}",
        ]
        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row, column=col, value=value)
            cell.alignment = center_align
            cell.border = thin_border

    summary_row = len(stats_list) + 2
    ws.cell(row=summary_row, column=1, value="Total").font = Font(bold=True)
    ws.cell(row=summary_row, column=3, value=sum(s.rag_on_total_refs for s in stats_list))
    ws.cell(row=summary_row, column=4, value=sum(s.rag_on_hallucinations for s in stats_list))
    ws.cell(row=summary_row, column=6, value=sum(s.rag_off_total_refs for s in stats_list))
    ws.cell(row=summary_row, column=7, value=sum(s.rag_off_hallucinations for s in stats_list))

    column_widths = [22, 14, 12, 14, 12, 13, 14, 12, 12]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width

    output_path = output_dir / "hallucination_summary.xlsx"
    wb.save(output_path)
    print(f"[Excel] Saved: {output_path}")
    return output_path


def generate_csv_summary(stats_list: List[HallucinationStats], output_dir: Path) -> Path:
    import csv
    output_path = output_dir / "hallucination_summary.csv"
    with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow([
            "Dimension", "Name", "RAG ON Total", "RAG ON Hallucinations", "RAG ON Rate",
            "RAG OFF Total", "RAG OFF Hallucinations", "RAG OFF Rate", "Rate Reduction"
        ])
        for stats in stats_list:
            writer.writerow([
                stats.dimension_key, stats.dimension_name,
                stats.rag_on_total_refs, stats.rag_on_hallucinations,
                f"{stats.rag_on_hallucination_rate:.1%}",
                stats.rag_off_total_refs, stats.rag_off_hallucinations,
                f"{stats.rag_off_hallucination_rate:.1%}",
                f"{stats.rate_reduction:.1%}",
            ])
    print(f"[CSV] Saved: {output_path}")
    return output_path


def generate_word_contrast(paragraphs: List[ContrastParagraph], output_dir: Path) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Warning: python-docx not installed, using Markdown fallback")
        return generate_markdown_contrast(paragraphs, output_dir)

    doc = Document()
    title = doc.add_heading("RAG Reference Contrast Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This document demonstrates the contrast between RAG ON and RAG OFF conditions. "
        "RAG ON provides precise regulation citations with specific clause numbers and technical indicators, "
        "while RAG OFF often produces generic statements without concrete basis."
    )

    doc.add_heading("Contrast Examples", level=1)
    for i, para in enumerate(paragraphs, 1):
        doc.add_heading(f"{i}. {para.dimension_name}", level=2)
        doc.add_heading("RAG OFF (Generic)", level=3)
        doc.add_paragraph(para.rag_off_text).italic = True
        doc.add_heading("RAG ON (Precise)", level=3)
        doc.add_paragraph(para.rag_on_text).bold = True
        doc.add_paragraph(f"Source: {para.regulation_source}")
        doc.add_paragraph(f"Explanation: {para.explanation}")
        doc.add_paragraph()

    output_path = output_dir / "contrast_analysis.docx"
    doc.save(output_path)
    print(f"[Word] Saved: {output_path}")
    return output_path


def generate_markdown_contrast(paragraphs: List[ContrastParagraph], output_dir: Path) -> Path:
    output_path = output_dir / "contrast_analysis.md"
    content = "# RAG Reference Contrast Analysis\n\n"
    content += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"

    for i, para in enumerate(paragraphs, 1):
        content += f"## {i}. {para.dimension_name}\n\n"
        content += f"**RAG OFF (Generic)**:\n\n> {para.rag_off_text}\n\n"
        content += f"**RAG ON (Precise)**:\n\n> {para.rag_on_text}\n\n"
        content += f"- Source: {para.regulation_source}\n"
        content += f"- Explanation: {para.explanation}\n\n---\n\n"

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"[Markdown] Saved: {output_path}")
    return output_path


def generate_kb_summary(docs: List[KnowledgeBaseDoc], output_dir: Path) -> Path:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
    except ImportError:
        return generate_kb_summary_csv(docs, output_dir)

    wb = Workbook()
    ws = wb.active
    ws.title = "KB Summary"

    header_fill = PatternFill(start_color="70AD47", end_color="70AD47", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")

    headers = ["Document", "Type", "Chunks", "Category"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    for row, doc in enumerate(docs, 2):
        ws.cell(row=row, column=1, value=doc.source)
        ws.cell(row=row, column=2, value=doc.doc_type)
        ws.cell(row=row, column=3, value=doc.chunk_count)
        ws.cell(row=row, column=4, value=doc.category)

    summary_row = len(docs) + 2
    ws.cell(row=summary_row, column=1, value="Total").font = Font(bold=True)
    ws.cell(row=summary_row, column=3, value=sum(d.chunk_count for d in docs))

    type_counts = {}
    for doc in docs:
        type_counts[doc.doc_type] = type_counts.get(doc.doc_type, 0) + 1

    row = summary_row + 2
    ws.cell(row=row, column=1, value="Type Summary").font = Font(bold=True)
    row += 1
    for doc_type, count in type_counts.items():
        ws.cell(row=row, column=1, value=doc_type)
        ws.cell(row=row, column=2, value=count)
        row += 1

    output_path = output_dir / "kb_summary.xlsx"
    wb.save(output_path)
    print(f"[KB Summary] Saved: {output_path}")
    return output_path


def generate_kb_summary_csv(docs: List[KnowledgeBaseDoc], output_dir: Path) -> Path:
    import csv
    output_path = output_dir / "kb_summary.csv"
    with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Document", "Type", "Chunks", "Category"])
        for doc in docs:
            writer.writerow([doc.source, doc.doc_type, doc.chunk_count, doc.category])
    print(f"[KB Summary CSV] Saved: {output_path}")
    return output_path


def get_mock_hallucination_stats() -> List[HallucinationStats]:
    return [
        HallucinationStats("land_use_planning", "土地利用规划", 12, 2, 0.167, 8, 5, 0.625, 0.458),
        HallucinationStats("infrastructure_planning", "基础设施规划", 10, 1, 0.10, 6, 4, 0.667, 0.567),
        HallucinationStats("ecological", "生态绿地规划", 15, 2, 0.133, 10, 7, 0.70, 0.567),
        HallucinationStats("disaster_prevention", "防震减灾规划", 18, 1, 0.056, 12, 9, 0.75, 0.694),
        HallucinationStats("heritage", "历史文保规划", 14, 2, 0.143, 9, 6, 0.667, 0.524),
    ]


def get_mock_contrast_paragraphs() -> List[ContrastParagraph]:
    return [
        ContrastParagraph(
            "disaster_prevention", "防震减灾规划",
            "规划区应避开地质灾害易发区域，确保居民安全。",
            "根据《广东省地质灾害防治条例》第十八条规定：崩塌隐患点避让距离不少于50米，滑坡隐患点避让距离不少于100米。规划区内5处崩塌隐患点应设置不少于50米的防护距离。",
            "《广东省地质灾害防治条例》第十八条",
            "RAG ON精准引用了具体条款编号和技术指标（50米、100米），而RAG OFF仅泛泛提及'避开灾害区域'。"
        ),
        ContrastParagraph(
            "heritage", "历史文保规划",
            "应保护历史文化资源，划定保护范围。",
            "依据《历史文化名城名镇名村保护条例》第二十七条：核心保护范围内不得进行新建、扩建活动。千年古檀树、古茶亭遗址应划为核心保护区，保护范围不少于周边20米。",
            "《历史文化名城名镇名村保护条例》第二十七条",
            "RAG ON准确引用了核心保护区的建设管控要求，并给出了具体的保护范围（20米）。"
        ),
        ContrastParagraph(
            "land_use_planning", "土地利用规划",
            "土地利用应符合相关规划要求，合理布局各类用地。",
            "根据《广东省村庄规划编制导则》第5.2条：村庄建设用地人均指标控制在100-150平方米。金田村现状人口500人，规划建设用地规模应控制在5-7.5公顷。",
            "《广东省村庄规划编制导则》第5.2条",
            "RAG ON引用了具体的技术指标（100-150平方米/人），并据此计算出村庄建设用地规模。"
        ),
        ContrastParagraph(
            "infrastructure_planning", "基础设施规划",
            "应完善基础设施建设，满足村民生活需求。",
            "依据《农村生活污水处理设施水污染物排放标准》（DB44/2209-2019）表1规定：出水排入环境水体时，CODcr≤60mg/L，氨氮≤8mg/L（水温>12℃）。",
            "DB44/2209-2019",
            "RAG ON精准引用了地方排放标准的具体编号和技术指标。"
        ),
        ContrastParagraph(
            "ecological", "生态绿地规划",
            "应加强生态保护，构建绿地系统。",
            "根据《广东省生态保护红线管理办法》第六条：生态保护红线内禁止工业化城镇化建设。金田村北部林地已划入生态保护红线，面积约8.5公顷，应严格保护。",
            "《广东省生态保护红线管理办法》第六条",
            "RAG ON引用了生态保护红线的管控要求，并给出了具体面积数据。"
        ),
    ]


def get_mock_kb_documents() -> List[KnowledgeBaseDoc]:
    return [
        KnowledgeBaseDoc("广东省地质灾害防治条例.pdf", "法规", 45, "laws"),
        KnowledgeBaseDoc("历史文化名城名镇名村保护条例.pdf", "法规", 38, "laws"),
        KnowledgeBaseDoc("广东省村庄规划编制导则.pdf", "技术标准", 52, "standards"),
        KnowledgeBaseDoc("农村生活污水处理设施水污染物排放标准.pdf", "技术标准", 28, "standards"),
        KnowledgeBaseDoc("广东省生态保护红线管理办法.pdf", "法规", 32, "laws"),
        KnowledgeBaseDoc("梅州市国土空间总体规划.pdf", "上位规划", 65, "plans"),
        KnowledgeBaseDoc("村庄规划编制技术指南.pdf", "技术标准", 42, "standards"),
        KnowledgeBaseDoc("农村人居环境整治技术规范.pdf", "技术标准", 35, "standards"),
    ]


def main(use_mock: bool = True):
    print("=" * 60)
    print("RAG Experiment Output Generator")
    print("=" * 60)

    ensure_rag_experiment_dirs()
    output_dir = RAG_HALLUCINATION_DIR / "outputs"
    output_dir.mkdir(parents=True, exist_ok=True)

    print("\n[1] Generating Excel summary...")
    stats_list = get_mock_hallucination_stats() if use_mock else []
    generate_excel_summary(stats_list, output_dir)

    print("\n[2] Generating Word contrast document...")
    paragraphs = get_mock_contrast_paragraphs() if use_mock else []
    generate_word_contrast(paragraphs, output_dir)

    print("\n[3] Generating Knowledge Base summary...")
    kb_docs = get_mock_kb_documents() if use_mock else []
    generate_kb_summary(kb_docs, output_dir)

    print("\n" + "=" * 60)
    print("All outputs generated!")
    print(f"Output directory: {output_dir}")
    print("=" * 60)


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--mock", action="store_true", help="Use mock data")
    args = parser.parse_args()
    main(use_mock=args.mock)

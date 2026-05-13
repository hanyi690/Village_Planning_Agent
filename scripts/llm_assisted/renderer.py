"""
Word Renderer V3.0

Renders generated content to Word documents using python-docx and docxtpl.
Handles template loading, text/table filling, and number verification.

V3.0 Changes:
- Template-driven rendering using docxtpl (Jinja2 in Word)
- Two-layer template architecture: shell template + content snippets
- Three-line table style support
"""

import json
import re
import sys
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, field
import logging

# Add project root to path for imports
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from scripts.llm_assisted.generator import GeneratedSection

logger = logging.getLogger(__name__)

# Template paths
TEMPLATE_DIR = project_root / "templates"
SHELL_TEMPLATE_PATH = TEMPLATE_DIR / "template_shell.docx"
SNIPPETS_DIR = TEMPLATE_DIR / "snippets"


@dataclass
class VerificationIssue:
    """Issue found during number verification"""
    field_name: str
    expected_value: Any
    actual_value: Any
    location: str  # Paragraph or table location
    severity: str  # "error", "warning", "info"


@dataclass
class RenderResult:
    """Result of Word document rendering"""
    output_path: str
    sections_rendered: int
    articles_count: int
    tables_filled: int
    verification_issues: List[VerificationIssue] = field(default_factory=list)
    success: bool = True
    error_message: str = ""


class WordRenderer:
    """
    Word Document Renderer V3.0

    Renders planning content to Word documents.
    Supports template-based rendering using docxtpl (Jinja2 in Word).

    Two-layer template architecture:
    - Shell template: 封面、目录、页眉页脚、章节标题样式
    - Content snippets: Jinja2 片段模板（表格、条文等）
    """

    def __init__(
        self,
        template_path: Optional[str] = None,
        verify_numbers: bool = True,
        use_template: bool = True
    ):
        """
        Initialize the renderer.

        Args:
            template_path: Path to Word template file (shell template)
            verify_numbers: Whether to verify numbers after rendering
            use_template: Whether to use template-driven rendering
        """
        self.template_path = template_path or str(SHELL_TEMPLATE_PATH)
        self.verify_numbers = verify_numbers
        self.use_template = use_template
        self._doc = None
        self._jinja_env = None

    def _get_jinja_env(self):
        """Lazy initialization of Jinja2 environment for snippets"""
        if self._jinja_env is None:
            from jinja2 import Environment, FileSystemLoader
            self._jinja_env = Environment(
                loader=FileSystemLoader(str(SNIPPETS_DIR)),
                autoescape=False
            )
        return self._jinja_env

    def render_with_template(
        self,
        context: Dict[str, Any],
        output_path: str
    ) -> RenderResult:
        """
        Render using docxtpl template (Jinja2 in Word).

        Args:
            context: Dictionary of variables for template rendering
            output_path: Output file path

        Returns:
            RenderResult
        """
        try:
            from docxtpl import DocxTemplate

            # Load template
            if Path(self.template_path).exists():
                doc = DocxTemplate(self.template_path)
            else:
                # Fallback to blank document
                logger.warning(f"Template not found: {self.template_path}, using blank document")
                return self._render_blank(context, output_path)

            # Render context
            doc.render(context)

            # Save
            doc.save(output_path)

            return RenderResult(
                output_path=output_path,
                sections_rendered=len(context.get("sections", [])),
                articles_count=context.get("article_count", 0),
                tables_filled=context.get("table_count", 0),
                success=True
            )

        except ImportError:
            logger.warning("docxtpl not available, falling back to blank document")
            return self._render_blank(context, output_path)
        except Exception as e:
            logger.error(f"Template rendering failed: {e}")
            return RenderResult(
                output_path=output_path,
                sections_rendered=0,
                articles_count=0,
                tables_filled=0,
                success=False,
                error_message=str(e)
            )

    def _render_blank(
        self,
        context: Dict[str, Any],
        output_path: str
    ) -> RenderResult:
        """Fallback: render to blank document"""
        sections = context.get("sections", [])
        project_name = context.get("project_name", "村庄规划")

        # Convert to GeneratedSection format and use existing render method
        from scripts.llm_assisted.generator import GeneratedSection

        generated_sections = []
        for section_data in sections:
            section = GeneratedSection(
                chapter_name=section_data.get("chapter_name", ""),
                dimension_key=section_data.get("dimension_key", ""),
                articles=section_data.get("articles", []),
                raw_response="",
                confidence=1.0
            )
            generated_sections.append(section)

        return self.render(
            sections=generated_sections,
            output_path=output_path,
            project_name=project_name,
            metadata=context.get("metadata")
        )

    def render(
        self,
        sections: List[GeneratedSection],
        output_path: str,
        project_name: str = "村庄规划",
        metadata: Optional[Dict[str, Any]] = None
    ) -> RenderResult:
        """
        Render sections to a Word document.

        Args:
            sections: List of GeneratedSection to render
            output_path: Output file path
            project_name: Project name for title
            metadata: Additional metadata (session_id, date, etc.)

        Returns:
            RenderResult with rendering details
        """
        try:
            # Create document
            if self.template_path and Path(self.template_path).exists():
                self._doc = self._load_template(self.template_path)
            else:
                self._doc = self._create_blank_document()

            # Add title and metadata
            self._add_header(project_name, metadata)

            # Render sections
            articles_count = 0
            for section in sections:
                articles_count += self._render_section(section)

            # Add verification summary
            issues = []
            if self.verify_numbers:
                issues = self._verify_document_numbers(sections)

            # Save document
            self._doc.save(output_path)

            return RenderResult(
                output_path=output_path,
                sections_rendered=len(sections),
                articles_count=articles_count,
                tables_filled=0,  # Will be updated if tables are used
                verification_issues=issues,
                success=True,
            )

        except Exception as e:
            logger.error(f"Rendering failed: {e}")
            return RenderResult(
                output_path=output_path,
                sections_rendered=0,
                articles_count=0,
                tables_filled=0,
                success=False,
                error_message=str(e),
            )

    def _load_template(self, template_path: str):
        """Load Word template"""
        try:
            from docx import Document
            return Document(template_path)
        except ImportError:
            raise ImportError("python-docx required for Word rendering")

    def _create_blank_document(self):
        """Create a blank Word document"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)

            return doc
        except ImportError:
            raise ImportError("python-docx required for Word rendering")

    def _add_header(
        self,
        project_name: str,
        metadata: Optional[Dict[str, Any]]
    ):
        """Add document header with title and metadata"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Title
        title = self._doc.add_heading(f"{project_name}规划文本", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if metadata:
            meta_para = self._doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if "session_id" in metadata:
                meta_para.add_run(f"会话ID: {metadata['session_id']}\n")
            if "date" in metadata:
                meta_para.add_run(f"生成日期: {metadata['date']}\n")

        self._doc.add_paragraph()  # Spacer

    def _render_section(self, section: GeneratedSection) -> int:
        """
        Render a single section to the document.

        Returns:
            Number of articles rendered
        """
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Section heading
        self._doc.add_heading(section.chapter_name, level=1)

        # Articles
        for article in section.articles:
            # 解析文章内容，分离表格和文本
            text_parts, tables = self._parse_markdown_content(article)

            # 添加文本部分
            for text in text_parts:
                para = self._doc.add_paragraph()
                run = para.add_run(text)
                run.font.size = Pt(12)

            # 添加表格
            for table_data in tables:
                self._add_table(table_data)

            # Add spacing after each article
            self._doc.add_paragraph()

        # Warnings if any
        if section.warnings:
            warn_para = self._doc.add_paragraph()
            warn_para.add_run("注意事项: ").bold = True
            for warning in section.warnings:
                warn_para.add_run(f"\n- {warning}")

        return len(section.articles)

    def _parse_markdown_content(self, content: str) -> Tuple[List[str], List[Dict]]:
        """
        解析Markdown内容，分离文本和表格

        Returns:
            (text_parts, tables) - 文本列表和表格数据列表
        """
        lines = content.split('\n')
        text_parts = []
        tables = []
        current_table = []
        in_table = False

        for line in lines:
            stripped = line.strip()

            # 检测表格行
            if stripped.startswith('|'):
                in_table = True
                current_table.append(line)
            else:
                if in_table and current_table:
                    # 表格结束，解析表格
                    table_data = self._parse_table_lines(current_table)
                    if table_data:
                        tables.append(table_data)
                    current_table = []
                    in_table = False

                # 非表格行，添加到文本
                if stripped:
                    text_parts.append(line)

        # 处理末尾表格
        if current_table:
            table_data = self._parse_table_lines(current_table)
            if table_data:
                tables.append(table_data)

        return text_parts, tables

    def _parse_table_lines(self, table_lines: List[str]) -> Optional[Dict]:
        """解析Markdown表格行"""
        if len(table_lines) < 2:
            return None

        # 解析表头
        header_line = table_lines[0]
        headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]

        # 跳过分隔符行
        data_lines = []
        for line in table_lines[1:]:
            stripped = line.strip()
            # 检查是否是分隔符行 (如 |---|---|)
            if re.match(r'^\|[\s\-:]+\|', stripped):
                continue
            data_lines.append(line)

        # 解析数据行
        rows = []
        for line in data_lines:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)

        return {
            'headers': headers,
            'rows': rows
        }

    def _add_table(self, table_data: Dict) -> None:
        """添加Word表格"""
        from docx.shared import Pt, Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT

        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])

        if not headers or not rows:
            return

        # 创建表格
        num_cols = len(headers)
        num_rows = len(rows) + 1  # 包含表头

        table = self._doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表头
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # 表头加粗
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)

        # 设置数据行
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)

        # 表格后添加空行
        self._doc.add_paragraph()

    def _add_three_line_table(self, table_data: Dict) -> None:
        """
        添加三线表格式的表格

        三线表特点：
        - 表头上下两条粗线
        - 表格底部一条粗线
        - 中间无竖线
        """
        from docx.shared import Pt, Inches, Twips
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])

        if not headers or not rows:
            return

        # 创建表格
        num_cols = len(headers)
        num_rows = len(rows) + 1  # 包含表头

        table = self._doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表头
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10.5)

        # 设置数据行
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10.5)

        # 设置三线表边框
        self._set_three_line_borders(table)

        # 表格后添加空行
        self._doc.add_paragraph()

    def _set_three_line_borders(self, table) -> None:
        """
        设置三线表边框样式

        - 表头顶部：1.5pt 粗线
        - 表头底部：1.5pt 粗线
        - 表格底部：1.5pt 粗线
        - 其他：无边框
        """
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )

        # 定义边框样式
        # 顶部粗线
        top_border = parse_xml(
            r'<w:top w:val="single" w:sz="15" w:space="0" w:color="000000" '
            r'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        # 底部粗线
        bottom_border = parse_xml(
            r'<w:bottom w:val="single" w:sz="15" w:space="0" w:color="000000" '
            r'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        # 内部水平线（细线）
        inside_h_border = parse_xml(
            r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000" '
            r'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        # 无垂直线
        inside_v_border = parse_xml(
            r'<w:insideV w:val="nil" '
            r'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        # 无左右边框
        left_border = parse_xml(
            r'<w:left w:val="nil" '
            r'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        right_border = parse_xml(
            r'<w:right w:val="nil" '
            r'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )

        # 创建 tblBorders 元素
        tblBorders = parse_xml(
            r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tblBorders.append(top_border)
        tblBorders.append(left_border)
        tblBorders.append(bottom_border)
        tblBorders.append(right_border)
        tblBorders.append(inside_h_border)
        tblBorders.append(inside_v_border)

        tblPr.append(tblBorders)

    def _verify_document_numbers(
        self,
        sections: List[GeneratedSection]
    ) -> List[VerificationIssue]:
        """
        Verify numbers in document match source data.

        Extracts numbers from rendered document and compares
        with source knowledge data.
        """
        issues = []

        # Extract all numbers from document
        doc_numbers = self._extract_numbers_from_doc()

        # Extract expected numbers from sections
        expected_numbers = self._extract_numbers_from_sections(sections)

        # Compare
        for field, expected in expected_numbers.items():
            if field in doc_numbers:
                actual = doc_numbers[field]
                if not self._numbers_match(expected, actual):
                    issues.append(VerificationIssue(
                        field_name=field,
                        expected_value=expected,
                        actual_value=actual,
                        location="Document text",
                        severity="warning"
                    ))
            else:
                issues.append(VerificationIssue(
                    field_name=field,
                    expected_value=expected,
                    actual_value=None,
                    location="Not found in document",
                    severity="info"
                ))

        return issues

    def _extract_numbers_from_doc(self) -> Dict[str, float]:
        """Extract all numbers from document paragraphs"""
        numbers = {}

        for para in self._doc.paragraphs:
            text = para.text

            # Find number patterns
            # Pattern: "XX公顷", "XX米", "XX%", etc.
            patterns = [
                r"([\d.]+)\s*公顷",
                r"([\d.]+)\s*米",
                r"([\d.]+)\s*公里",
                r"([\d.]+)\s*%",
                r"([\d.]+)\s*平方米",
            ]

            for pattern in patterns:
                matches = re.findall(pattern, text)
                for match in matches:
                    try:
                        value = float(match)
                        # Use surrounding text as field name
                        context = re.search(r"([^\d]+)" + pattern, text)
                        if context:
                            field = context.group(1).strip()
                            numbers[field] = value
                    except ValueError:
                        continue

        return numbers

    def _extract_numbers_from_sections(
        self,
        sections: List[GeneratedSection]
    ) -> Dict[str, float]:
        """Extract expected numbers from section data"""
        numbers = {}

        for section in sections:
            # This would need the original coordinated data
            # For now, extract from article text
            for article in section.articles:
                patterns = [
                    r"([\d.]+)\s*公顷",
                    r"([\d.]+)\s*米",
                    r"([\d.]+)\s*公里",
                    r"([\d.]+)\s*%",
                ]

                for pattern in patterns:
                    matches = re.findall(pattern, article)
                    for match in matches:
                        try:
                            numbers[f"{section.dimension_key}_{pattern}"] = float(match)
                        except ValueError:
                            continue

        return numbers

    def _numbers_match(self, expected: Any, actual: Any) -> bool:
        """Compare two numbers with tolerance"""
        try:
            exp = float(expected)
            act = float(actual)
            return abs(exp - act) < 0.01  # Allow small rounding differences
        except (TypeError, ValueError):
            return False

    def fill_table(
        self,
        table_index: int,
        data: List[Dict[str, Any]],
        headers: Optional[List[str]] = None
    ) -> int:
        """
        Fill a table in the document with data.

        Args:
            table_index: Index of table in document
            data: List of row data dicts
            headers: Optional header row

        Returns:
            Number of rows filled
        """
        if table_index >= len(self._doc.tables):
            logger.warning(f"Table {table_index} not found in document")
            return 0

        table = self._doc.tables[table_index]
        rows_filled = 0

        # Add header if provided
        if headers and len(table.rows) > 0:
            row = table.rows[0]
            for i, header in enumerate(headers):
                if i < len(row.cells):
                    row.cells[i].text = header

        # Add data rows
        start_row = 1 if headers else 0
        for i, row_data in enumerate(data):
            row_idx = start_row + i

            # Add new row if needed
            if row_idx >= len(table.rows):
                table.add_row()

            row = table.rows[row_idx]
            for j, (key, value) in enumerate(row_data.items()):
                if j < len(row.cells):
                    row.cells[j].text = str(value)

            rows_filled += 1

        return rows_filled

    def generate_verification_report(
        self,
        issues: List[VerificationIssue],
        output_path: str
    ) -> None:
        """Generate a verification report as Markdown"""
        lines = ["# 数字核验报告", ""]

        errors = [i for i in issues if i.severity == "error"]
        warnings = [i for i in issues if i.severity == "warning"]
        infos = [i for i in issues if i.severity == "info"]

        if errors:
            lines.append("## 错误")
            for issue in errors:
                lines.append(f"- **{issue.field_name}**: 期望 {issue.expected_value}, 实际 {issue.actual_value}")
            lines.append("")

        if warnings:
            lines.append("## 警告")
            for issue in warnings:
                lines.append(f"- **{issue.field_name}**: 期望 {issue.expected_value}, 实际 {issue.actual_value}")
            lines.append("")

        if infos:
            lines.append("## 信息")
            for issue in infos:
                lines.append(f"- **{issue.field_name}**: 期望 {issue.expected_value}, 未在文档中找到")
            lines.append("")

        if not issues:
            lines.append("*所有数字核验通过*")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))


def render_to_word(
    sections: List[GeneratedSection],
    output_path: str,
    project_name: str = "村庄规划",
    template_path: Optional[str] = None
) -> RenderResult:
    """
    Convenience function to render sections to Word.

    Args:
        sections: List of GeneratedSection
        output_path: Output file path
        project_name: Project name
        template_path: Optional template path

    Returns:
        RenderResult
    """
    renderer = WordRenderer(template_path=template_path)
    return renderer.render(sections, output_path, project_name)


__all__ = [
    "WordRenderer",
    "RenderResult",
    "VerificationIssue",
    "render_to_word",
]
"""
Layer report exporter.

Exports layer-specific reports (L1/现状分析, L2/规划思路, L3/详细规划)
in MD/JSON/DOCX formats with optional RAG dimension-level details.
"""

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from dataclasses import dataclass, field

from app.services.report_store import ReportStore

logger = logging.getLogger(__name__)

# Layer names
LAYER_NAMES = {
    1: "现状分析",
    2: "规划思路",
    3: "详细规划",
}

_TYPE_LABELS = {
    "policy": "法规",
    "standard": "标准",
    "textbook": "教材",
    "case": "案例",
    "planning": "上位规划",
    "guide": "指南",
}


@dataclass
class LayerExportConfig:
    """Configuration for layer report export."""
    session_id: str
    output_dir: str = "output/planning_text"
    output_formats: Tuple[str, ...] = ("md", "json")
    include_rag_details: bool = False
    include_summaries: bool = True


@dataclass
class DimensionReport:
    """Single dimension report with metadata."""
    dimension_key: str
    dimension_name: str
    layer: int
    content: str
    summary: Optional[Dict[str, Any]] = None
    rag_sources: List[Dict[str, Any]] = field(default_factory=list)


@dataclass
class LayerExportResult:
    """Result of layer export operation."""
    layer: int
    layer_name: str
    dimensions: List[DimensionReport]
    md_path: Optional[str] = None
    json_path: Optional[str] = None
    docx_path: Optional[str] = None
    errors: List[str] = field(default_factory=list)


class LayerExporter:
    """Exports layer-specific reports in multiple formats."""

    def __init__(self, config: LayerExportConfig):
        self.config = config
        self.store = ReportStore.get_instance()

    async def export_layer(self, layer: int) -> LayerExportResult:
        """Export a single layer report.

        Args:
            layer: Layer number (1, 2, or 3)

        Returns:
            LayerExportResult with paths to exported files
        """
        result = LayerExportResult(
            layer=layer,
            layer_name=LAYER_NAMES.get(layer, f"Layer {layer}"),
            dimensions=[],
        )

        # Fetch reports
        reports = await self.store.get_layer_reports(
            self.config.session_id, layer
        )
        if not reports:
            result.errors.append(f"Layer {layer} 无数据")
            return result

        # Fetch summaries if needed
        summaries = {}
        if self.config.include_summaries:
            try:
                summaries = await self.store.get_layer_summaries(
                    self.config.session_id, layer
                )
            except Exception as e:
                logger.warning(f"[LayerExporter] Failed to fetch summaries: {e}")

        # Fetch RAG sources if needed
        rag_sources = {}
        if self.config.include_rag_details:
            rag_sources = await self._fetch_rag_sources(layer)

        # Build dimension reports
        dim_names = self._get_dimension_names()
        for dim_key, content in reports.items():
            dim_report = DimensionReport(
                dimension_key=dim_key,
                dimension_name=dim_names.get(dim_key, dim_key),
                layer=layer,
                content=content,
                summary=summaries.get(dim_key),
                rag_sources=rag_sources.get(dim_key, []),
            )
            result.dimensions.append(dim_report)

        # Export files
        out_dir = Path(self.config.output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)

        if "md" in self.config.output_formats:
            result.md_path = self._export_md(result, out_dir)

        if "json" in self.config.output_formats:
            result.json_path = self._export_json(result, out_dir)

        if "docx" in self.config.output_formats:
            result.docx_path = self._export_docx(result, out_dir)

        return result

    async def export_all_layers(self) -> Dict[int, LayerExportResult]:
        """Export all three layer reports.

        Returns:
            Dict mapping layer number to LayerExportResult
        """
        import asyncio
        results = {}
        for layer in [1, 2, 3]:
            results[layer] = await self.export_layer(layer)
        return results

    # ---- Private methods ----

    async def _fetch_rag_sources(self, layer: int) -> Dict[str, List[Dict]]:
        """Fetch RAG knowledge sources for a layer."""
        rag_map = {}
        try:
            reports = await self.store.get_layer_reports_with_sources(
                self.config.session_id, layer
            )
            for dim_key, data in reports.items():
                ks = data.get("knowledge_sources")
                if not ks:
                    continue
                chunks = self._extract_chunks(ks)
                if chunks:
                    rag_map[dim_key] = self._process_rag_chunks(chunks)
        except Exception as e:
            logger.warning(f"[LayerExporter] Failed to fetch RAG sources: {e}")
        return rag_map

    def _process_rag_chunks(self, chunks: List[Dict]) -> List[Dict[str, Any]]:
        """Process RAG chunks into structured format."""
        seen = set()
        result = []
        for c in chunks:
            title = c.get("title", "")
            doc_type = c.get("doc_type", "")
            source = c.get("source", "")
            if not title and source:
                title, doc_type = self._lookup_source(source)
            if not title and source:
                title = source.rsplit("_minerU_parsed", 1)[0].removesuffix(".md").strip()
            if title and title not in seen:
                seen.add(title)
                result.append({
                    "title": title,
                    "doc_type": doc_type,
                    "label": _TYPE_LABELS.get(doc_type, doc_type),
                    "score": c.get("score", 0),
                    "content_preview": c.get("content", "")[:200] if c.get("content") else "",
                })
        return result

    def _export_md(self, result: LayerExportResult, out_dir: Path) -> str:
        """Export layer report as markdown."""
        lines = [
            f"# {result.layer_name}报告",
            "",
            f"**项目**: {self.config.session_id}",
            f"**层级**: Layer {result.layer} - {result.layer_name}",
            f"**维度数**: {len(result.dimensions)}",
            f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "---",
            "",
        ]

        for dim in result.dimensions:
            lines.append(f"## {dim.dimension_name}")
            lines.append("")
            lines.append(dim.content)
            lines.append("")

            # Add RAG details if enabled
            if self.config.include_rag_details and dim.rag_sources:
                lines.append("### 参考知识来源")
                lines.append("")
                for src in dim.rag_sources:
                    lines.append(f"- [{src['label']}] {src['title']}")
                lines.append("")

            # Add summary if available
            if self.config.include_summaries and dim.summary:
                lines.append("### 结构化摘要")
                lines.append("")
                lines.append("```json")
                lines.append(json.dumps(dim.summary, ensure_ascii=False, indent=2))
                lines.append("```")
                lines.append("")

            lines.append("---")
            lines.append("")

        md_content = "\n".join(lines)
        md_path = out_dir / f"L{result.layer}_{result.layer_name}.md"
        md_path.write_text(md_content, encoding='utf-8')
        logger.info(f"[LayerExporter] MD saved: {md_path}")
        return str(md_path)

    def _export_json(self, result: LayerExportResult, out_dir: Path) -> str:
        """Export layer report as JSON."""
        data = {
            "meta": {
                "session_id": self.config.session_id,
                "layer": result.layer,
                "layer_name": result.layer_name,
                "dimension_count": len(result.dimensions),
                "generated_at": datetime.now().isoformat(),
                "include_rag_details": self.config.include_rag_details,
                "include_summaries": self.config.include_summaries,
            },
            "dimensions": [],
        }

        for dim in result.dimensions:
            dim_data = {
                "key": dim.dimension_key,
                "name": dim.dimension_name,
                "layer": dim.layer,
                "content": dim.content,
                "word_count": len(dim.content) if dim.content else 0,
            }
            if self.config.include_summaries and dim.summary:
                dim_data["summary"] = dim.summary
            if self.config.include_rag_details and dim.rag_sources:
                dim_data["rag_sources"] = dim.rag_sources
            data["dimensions"].append(dim_data)

        json_path = out_dir / f"L{result.layer}_{result.layer_name}.json"
        json_path.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding='utf-8',
        )
        logger.info(f"[LayerExporter] JSON saved: {json_path}")
        return str(json_path)

    def _export_docx(self, result: LayerExportResult, out_dir: Path) -> Optional[str]:
        """Export layer report as DOCX (optional)."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading(f"{result.layer_name}报告", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            doc.add_paragraph(f"项目: {self.config.session_id}")
            doc.add_paragraph(f"层级: Layer {result.layer} - {result.layer_name}")
            doc.add_paragraph(f"维度数: {len(result.dimensions)}")
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")

            for dim in result.dimensions:
                doc.add_heading(dim.dimension_name, level=2)

                # Split content into paragraphs
                if dim.content:
                    for para in dim.content.split('\n\n'):
                        para = para.strip()
                        if para:
                            doc.add_paragraph(para)

                # RAG details
                if self.config.include_rag_details and dim.rag_sources:
                    doc.add_heading("参考知识来源", level=3)
                    for src in dim.rag_sources:
                        doc.add_paragraph(
                            f"[{src['label']}] {src['title']}",
                            style='List Bullet',
                        )

                # Summary
                if self.config.include_summaries and dim.summary:
                    doc.add_heading("结构化摘要", level=3)
                    doc.add_paragraph(
                        json.dumps(dim.summary, ensure_ascii=False, indent=2)
                    )

            docx_path = out_dir / f"L{result.layer}_{result.layer_name}.docx"
            doc.save(str(docx_path))
            logger.info(f"[LayerExporter] DOCX saved: {docx_path}")
            return str(docx_path)

        except ImportError:
            logger.warning("[LayerExporter] python-docx not installed, skipping DOCX export")
            return None
        except Exception as e:
            logger.warning(f"[LayerExporter] DOCX export failed: {e}")
            return None

    def _get_dimension_names(self) -> Dict[str, str]:
        """Get dimension names from config."""
        try:
            from app.config.loader import load_phases_config
            phases = load_phases_config()
            dim_names = {}
            for phase in phases.get("phases", []):
                for dim in phase.get("dimensions", []):
                    dim_names[dim["key"]] = dim.get("name", dim["key"])
            return dim_names
        except Exception:
            return {}

    @staticmethod
    def _extract_chunks(ks: Any) -> List[Dict]:
        """Extract chunks from knowledge sources."""
        if isinstance(ks, dict):
            chunks = ks.get("retrieved_chunks", [])
            if isinstance(chunks, list):
                return chunks
        if isinstance(ks, list):
            return [c for c in ks if isinstance(c, dict)]
        return []

    @staticmethod
    def _lookup_source(source: str) -> Tuple[str, str]:
        """Lookup source metadata from cache."""
        try:
            cache_path = (
                Path(__file__).parent.parent.parent.parent.parent.parent
                / "data" / "RAG_doc" / "_cache" / "source_metadata_map.json"
            )
            if not hasattr(LayerExporter, '_source_map'):
                LayerExporter._source_map = {}
                if cache_path.exists():
                    LayerExporter._source_map = json.loads(
                        cache_path.read_text("utf-8")
                    )
            info = LayerExporter._source_map.get(source, {})
            return info.get("title", ""), info.get("doc_type", "")
        except Exception:
            return "", ""


__all__ = [
    "LayerExportConfig",
    "LayerExporter",
    "LayerExportResult",
    "DimensionReport",
    "LAYER_NAMES",
]

"""
Document Loader - 统一文档加载器

Supports: doc, docx, pdf, pptx, xlsx, epub, html, txt, md
Architecture: MarkItDown (primary) + fallback chain per format
"""

import logging
import re
import tempfile
import subprocess
from pathlib import Path
from typing import Any, Dict, List, Optional, Callable

logger = logging.getLogger(__name__)


# ==========================================
# 文件类型检测
# ==========================================

_MIME_MAP = {
    'application/pdf': 'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/msword': 'doc',
    'application/vnd.openxmlformats-officedocument.presentationml.presentation': 'pptx',
    'application/vnd.ms-powerpoint': 'ppt',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
    'application/vnd.ms-excel': 'xls',
    'application/epub+zip': 'epub',
    'text/html': 'html',
    'text/plain': 'txt',
    'text/markdown': 'md',
    'application/octet-stream': None,
}

_EXT_MAP = {
    '.pdf': 'pdf',
    '.docx': 'docx',
    '.doc': 'doc',
    '.pptx': 'pptx',
    '.ppt': 'ppt',
    '.xlsx': 'xlsx',
    '.xls': 'xls',
    '.epub': 'epub',
    '.html': 'html',
    '.htm': 'html',
    '.txt': 'txt',
    '.md': 'md',
    '.markdown': 'md',
}


def classify_file_type(filename: str) -> str:
    """Classify uploaded file type for API layer use

    Args:
        filename: File name string

    Returns:
        File type: "document" | "geojson" | "shapefile" | "kml" | "gis_file"
    """
    name_lower = filename.lower()
    if name_lower.endswith('.geojson'):
        return 'geojson'
    if name_lower.endswith('.zip') or name_lower.endswith('.shp'):
        return 'shapefile'
    if name_lower.endswith('.kml') or name_lower.endswith('.kmz'):
        return 'kml'
    if any(name_lower.endswith(ext) for ext in ['.gpkg', '.gpx', '.gml']):
        return 'gis_file'
    return 'document'


class FileTypeDetector:
    """Detect real file type via MIME and extension"""

    @staticmethod
    def detect(file_path: Path) -> str:
        import mimetypes
        from pathlib import Path

        file_path = Path(file_path)

        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type and mime_type in _MIME_MAP and _MIME_MAP[mime_type]:
            return _MIME_MAP[mime_type]

        ext = file_path.suffix.lower()
        if ext in _EXT_MAP:
            return _EXT_MAP[ext]

        if ext in ('.zip', '.shp'):
            return 'shapefile'
        if ext in ('.geojson',):
            return 'geojson'
        if ext in ('.kml', '.kmz'):
            return 'kml'

        return ext.lstrip('.') if ext else 'unknown'


# ==========================================
# Markdown Cleaner
# ==========================================

class MarkdownCleaner:
    """Clean Markdown content: remove excess blank lines, fix encoding"""

    @staticmethod
    def clean(content: str) -> str:
        content = re.sub(r'\n{3,}', '\n\n', content)
        content = re.sub(r'[ \t]+$', '', content, flags=re.MULTILINE)
        content = content.replace('\r\n', '\n')
        content = content.replace('\r', '\n')
        if content.startswith('\ufeff'):
            content = content[1:]
        content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', content)
        return content.strip()


# ==========================================
# .doc Converter
# ==========================================

class DocToDocxConverter:
    """Convert .doc (OLE) files to .docx using LibreOffice"""

    @staticmethod
    def convert(doc_path: Path) -> Optional[Path]:
        import shutil

        if not doc_path.exists():
            logger.warning(f"[DocToDocxConverter] File not found: {doc_path}")
            return None

        libreoffice = shutil.which('libreoffice') or shutil.which('soffice')
        if not libreoffice:
            logger.warning("[DocToDocxConverter] LibreOffice not installed")
            return None

        try:
            output_dir = tempfile.mkdtemp(prefix='doc_convert_')
            cmd = [
                libreoffice,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', output_dir,
                str(doc_path),
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            if result.returncode != 0:
                logger.error(f"[DocToDocxConverter] Conversion failed: {result.stderr}")
                return None

            output_name = doc_path.stem + '.docx'
            output_path = Path(output_dir) / output_name
            if output_path.exists():
                return output_path

            for f in Path(output_dir).glob('*.docx'):
                return f

            logger.warning("[DocToDocxConverter] Output file not found")
            return None

        except subprocess.TimeoutExpired:
            logger.error("[DocToDocxConverter] Conversion timed out")
            return None
        except Exception as e:
            logger.error(f"[DocToDocxConverter] Conversion error: {e}")
            return None


# ==========================================
# MarkItDown Unified Loader
# ==========================================

class MarkItDownLoader:
    """Unified document loader using MarkItDown + fallback chain"""

    def __init__(self, file_path: Path, category: Optional[str] = None):
        self.file_path = Path(file_path)
        self.file_type = FileTypeDetector.detect(self.file_path)
        self.category = category

    def load(self) -> List:
        from langchain_core.documents import Document

        content = self._extract_content()

        if not content or len(content.strip()) < 10:
            logger.warning(f"[MarkItDownLoader] Content empty or too short: {self.file_path}")
            return []

        content = MarkdownCleaner.clean(content)

        doc = Document(
            page_content=content,
            metadata={
                'source': self.file_path.name,
                'type': self.file_type,
                'file_path': str(self.file_path),
                'category': self.category or 'unknown',
            }
        )
        return [doc]

    def _extract_content(self) -> str:
        content = self._try_markitdown()
        if content:
            return content

        fallback = self._get_fallback()
        if fallback:
            return fallback()

        return self._try_plain_text()

    def _try_markitdown(self) -> Optional[str]:
        try:
            from markitdown import MarkItDown
            md = MarkItDown()
            result = md.convert(str(self.file_path))
            if result and result.text_content:
                return result.text_content
        except ImportError:
            logger.debug("[MarkItDownLoader] markitdown not installed")
        except Exception as e:
            logger.warning(f"[MarkItDownLoader] markitdown parse failed: {e}")
        return None

    def _get_fallback(self) -> Optional[Callable]:
        if self.file_type == 'txt':
            return self._try_plain_text
        elif self.file_type == 'pdf':
            return self._try_pdf_fallback
        elif self.file_type == 'docx':
            return self._try_docx_fallback
        elif self.file_type == 'doc':
            return self._try_doc_fallback
        elif self.file_type in ('html', 'htm'):
            return self._try_html_fallback
        return None

    def _try_pdf_fallback(self) -> Optional[str]:
        try:
            from app.utils.pdf_fallback import parse_pdf_with_fallback
            return parse_pdf_with_fallback(self.file_path)
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"[MarkItDownLoader] PDF fallback failed: {e}")
        return None

    def _try_docx_fallback(self) -> str:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(self.file_path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    paragraphs.append(' | '.join(cells))
            return '\n\n'.join(paragraphs)
        except ImportError:
            logger.debug("[MarkItDownLoader] python-docx not installed")
        except Exception as e:
            logger.warning(f"[MarkItDownLoader] docx fallback failed: {e}")
        return self._try_plain_text()

    def _try_doc_fallback(self) -> Optional[str]:
        converter = DocToDocxConverter()
        docx_path = converter.convert(self.file_path)
        if docx_path:
            try:
                self.file_path = docx_path
                self.file_type = 'docx'
                return self._try_docx_fallback()
            finally:
                pass
        return None

    def _try_html_fallback(self) -> str:
        try:
            from bs4 import BeautifulSoup
            html = self.file_path.read_text(encoding='utf-8', errors='replace')
            soup = BeautifulSoup(html, 'html.parser')
            for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
                tag.decompose()
            return soup.get_text('\n', strip=True)
        except ImportError:
            logger.debug("[MarkItDownLoader] beautifulsoup4 not installed")
        except Exception as e:
            logger.warning(f"[MarkItDownLoader] html fallback failed: {e}")
        return self._try_plain_text()

    def _try_plain_text(self) -> str:
        for encoding in ['utf-8', 'gbk', 'gb18030', 'latin-1']:
            try:
                return self.file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return self.file_path.read_text(encoding='utf-8', errors='replace')


# ==========================================
# Factory
# ==========================================

def _create_loader(
    file_path: Path,
    real_type: str,
    category: Optional[str] = None,
):
    supported = {'pdf', 'docx', 'doc', 'pptx', 'xlsx', 'epub', 'html', 'txt', 'md'}
    if real_type in supported:
        return MarkItDownLoader(file_path, category=category)
    return None


# ==========================================
# Batch Directory Loader
# ==========================================

def load_documents_from_directory(
    directory: Path,
    category_map: Optional[Dict[str, str]] = None,
    progress_callback: Optional[Callable[[float, str], None]] = None,
) -> List:
    from pathlib import Path

    directory = Path(directory)
    if not directory.exists():
        logger.warning(f"[load_documents] Directory not found: {directory}")
        return []

    supported_extensions = {'.pdf', '.docx', '.doc', '.pptx', '.xlsx',
                            '.epub', '.html', '.htm', '.txt', '.md'}
    files = [f for f in directory.rglob('*') if f.suffix.lower() in supported_extensions]

    if not files:
        logger.warning(f"[load_documents] No supported documents in: {directory}")
        return []

    all_docs = []
    total = len(files)

    for idx, file_path in enumerate(files):
        if progress_callback:
            progress = (idx / total) * 100 if total > 0 else 0
            progress_callback(progress, f"Loading: {file_path.name}")

        try:
            real_type = FileTypeDetector.detect(file_path)
            category = (category_map or {}).get(file_path.name)
            loader = _create_loader(file_path, real_type, category=category)
            if loader:
                docs = loader.load()
                all_docs.extend(docs)
                logger.info(f"[load_documents] Loaded {file_path.name} ({len(docs)} chunks)")
        except Exception as e:
            logger.error(f"[load_documents] Failed to load {file_path.name}: {e}")

    if progress_callback:
        progress_callback(100.0, f"Complete: {len(all_docs)} chunks from {total} files")

    return all_docs


__all__ = [
    "FileTypeDetector",
    "MarkItDownLoader",
    "MarkdownCleaner",
    "DocToDocxConverter",
    "classify_file_type",
    "load_documents_from_directory",
    "_create_loader",
]
